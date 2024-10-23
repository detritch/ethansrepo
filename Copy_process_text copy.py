import re
import os
from docx import Document
from docx.shared import Pt

def process_text(input_text, document):
    lines = input_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if line.startswith('###') or re.match(r'^\d+\.', line):
            # Main header (remove '###' and '**') or numbered item
            formatted_line = line.lstrip('#').strip().strip('**')
            # Remove leading number and dot if present
            formatted_line = re.sub(r'^\d+\.\s*', '', formatted_line)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(formatted_line)
            run.font.size = Pt(22)
            run.bold = True
        elif line.startswith('-'):
            # Subheader (remove '-' and '**')
            formatted_line = line.lstrip('-').strip().strip('**')
            paragraph = document.add_paragraph()
            run = paragraph.add_run(formatted_line)
            run.font.size = Pt(11)
            run.bold = True
        elif line.startswith('     -'):
            # Indented bullet point (remove '     -')
            formatted_line = line.lstrip(' -').strip()
            paragraph = document.add_paragraph(formatted_line, style='List Bullet')
        else:
            # Normal text
            paragraph = document.add_paragraph(line)

def get_first_three_words(text):
    words = re.findall(r'\w+', text)
    return '_'.join(words[:3]).lower()

def process_files(input_folder, output_folder):
    # Ensure the output directory exists
    os.makedirs(output_folder, exist_ok=True)
    
    # Process all text files in the input folder
    for filename in os.listdir(input_folder):
        if filename.endswith('.txt'):
            input_path = os.path.join(input_folder, filename)
            base_output_name = f"processed_{os.path.splitext(filename)[0]}.docx"
            output_path = os.path.join(output_folder, base_output_name)
            
            with open(input_path, 'r', encoding='utf-8') as input_file:
                input_text = input_file.read()
            
            document = Document()
            process_text(input_text, document)
            
            # Check if the file already exists
            if os.path.exists(output_path):
                # Use the first three words as the filename
                first_three_words = get_first_three_words(input_text)
                new_filename = f"{first_three_words}.docx"
                output_path = os.path.join(output_folder, new_filename)
                
                # If even this filename exists, append a number
                counter = 1
                while os.path.exists(output_path):
                    new_filename = f"{first_three_words}_{counter}.docx"
                    output_path = os.path.join(output_folder, new_filename)
                    counter += 1
            
            document.save(output_path)
            
            print(f"Processed {filename} and saved as {os.path.basename(output_path)}")

if __name__ == "__main__":
    input_folder = r"F:\October 2024 coding\Input Text"
    output_folder = r"F:\October 2024 coding\Output Text"
    
    process_files(input_folder, output_folder)
    
    print("All files have been processed.")
